import docx

# import nightsawayforms codelocal files
import campdeets
import configuration
import nightsawayforms


def nan_form(directory, leader, myGroup, event):
    """Build the Nights Away Notification (NAN) form for the camp"""
    # create new document object
    doc = docx.Document()
    print("Generating NAN form...")

    # add heading and introductory paragraphs
    h = doc.add_paragraph()
    head = h.add_run('Nights Away Notification')
    head.bold = True
    head.italic = True
    head.font.size = docx.shared.Pt(28)
    head.font.color.rgb = docx.shared.RGBColor(111, 111, 111)

    p = doc.add_paragraph()
    p.add_run("This form provides the information a Commissioner requires to ")
    r = p.add_run("APPROVE")
    r.bold = True
    p.add_run(" an event to take place (i.e. POR 9.1b/9.1c). " \
                + "The Permit holder is responsible for ensuring that the appropriate Commissioner is informed " \
                + "about each Colony, Pack, Troop, or Unit attending a nights away event (even a District or County event). " \
                + "This can also be done online in the event function of Compass.")

    p = doc.add_paragraph()
    p.add_run("For all Nights Away experiences all of the information below should be with your District Commissioner (or appointee) ")
    r = p.add_run("SEVEN")
    r.bold = True
    p.add_run(" days before the event (in normal circumstances). " \
                + "How the information is passed on will depend on local arrangements " \
                + "(this may be for example by telephone call, e-mail or fax).")

    # add main table for the programme
    form = doc.add_table(rows=13, cols=4)
    form.style = 'Table Grid'
    form.style.font.size = docx.shared.Pt(8)
    
    # add leader details to form
    cell = form.cell(0, 0)
    cell.width = docx.shared.Cm(4)
    cell.paragraphs[0].text = "Permit Holder's Name"
    shading = docx.oxml.parse_xml(r'<w:shd {} w:fill="B2B2B2"/>'.format(docx.oxml.ns.nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading)

    cell = form.cell(0, 1)
    cell.width = docx.shared.Cm(5)
    cell.paragraphs[0].text = leader.name

    cell = form.cell(0, 2)
    cell.width = docx.shared.Cm(4)
    cell.paragraphs[0].text = "Telephone"
    shading = docx.oxml.parse_xml(r'<w:shd {} w:fill="B2B2B2"/>'.format(docx.oxml.ns.nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading)

    cell = form.cell(0, 3)
    cell.width = docx.shared.Cm(5)
    cell.paragraphs[0].text = leader.telephone

    cell = form.cell(1, 0)
    cell.paragraphs[0].text = "Membership #"
    shading = docx.oxml.parse_xml(r'<w:shd {} w:fill="B2B2B2"/>'.format(docx.oxml.ns.nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading)

    cell = form.cell(1, 1)
    cell.paragraphs[0].text = leader.memberNo

    cell = form.cell(1, 2)
    cell.paragraphs[0].text = "Email"
    shading = docx.oxml.parse_xml(r'<w:shd {} w:fill="B2B2B2"/>'.format(docx.oxml.ns.nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading)

    cell = form.cell(1, 3)
    cell.paragraphs[0].text = leader.email

    cell = form.cell(2, 0)
    cell.paragraphs[0].text = "Group and District"
    shading = docx.oxml.parse_xml(r'<w:shd {} w:fill="B2B2B2"/>'.format(docx.oxml.ns.nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading)

    cell = form.cell(2, 1)
    cell.paragraphs[0].text = myGroup.group + " / " + myGroup.district

    cell = form.cell(2, 2)
    cell.paragraphs[0].text = "Section"
    shading = docx.oxml.parse_xml(r'<w:shd {} w:fill="B2B2B2"/>'.format(docx.oxml.ns.nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading)

    cell = form.cell(2, 3)
    cell.paragraphs[0].text = myGroup.section



    # save document
    try:
        #docName = 'NAN form - ' + event.camp + ' ' + event.campYear + '.docx'
        docName = "NAN form - Summer camp 2018.docx"
        print("Saving:", docName)
        doc.save(directory + docName)
    except:
        print("Failed to save " + docName)


if __name__ == '__main__':
    event = campdeets.campdeets("summer camp", "The Hardwick Estate, Pangbourne", "11-08-2018", "18-08-2018", True, True, True)
    myGroup = configuration.config_reader()
    leader = campdeets.campLeader("Ace Broderick", "ASL", "123456", "0123456789", "ace@broderick.com")
    directory = nightsawayforms.camp_directory()
    nan_form(directory, leader, myGroup, event)