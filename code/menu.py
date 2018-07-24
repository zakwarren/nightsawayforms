import docx


def menu(doc, docName, directory, event):
    """Build the menu for the camp"""
    print("Generating menu...")
    # add heading
    doc.add_heading('Menu', 4)

    # add main table for the menu
    rowCount = event.nightsAway + 2
    prog = doc.add_table(rows=rowCount, cols=2)
    prog.style = 'Table Grid'
    # add headings
    cell = prog.cell(0, 0)
    cell.width = docx.shared.Cm(3)
    p = cell.add_paragraph()
    r = p.add_run('Day')
    r.bold = True
    cell = prog.cell(0, 1)
    cell.width = docx.shared.Cm(20)
    p = cell.add_paragraph()
    r = p.add_run('Meals')
    r.bold = True

    # populate main table for the menu
    daysAway = event.nightsAway + 1
    i = 1
    while i <= daysAway:
        # add day
        cellD = prog.cell(i, 0)
        d = i - 1
        day = event.daysCamping[d]
        cellD.text = day
        # add menu details
        cellA = prog.cell(i, 1)
        a = 1
        while True:
            act = input("How many meals will there be on " + day + "? ")
            # check is a number
            try:
                act = int(act)
                break
            except:
                print("That isn't a number!")
        # check if less than three meals a day are selected
        while act < 3:
            print("You've only selected " + str(act) + " meals on " + day + "!")
            check = input("Is this correct (y/n)? ")
            if check != 'y':
                act = input("How many meals will there be on " + day + "? ")
            elif check == 'y':
                break
        # get meals and append to menu
        while a <= int(act):
            activity = input(day + "'s meal " + str(a) + ": ")
            if a == 1:
                cellA.text = activity
            else:
                pA = cellA.add_paragraph()
                pA.add_run(activity)
            a += 1
        # iterate through loop
        i += 1
    
    # save document
    try:
        docName = 'Menu' + docName
        print("Saving:", docName)
        doc.save(directory + docName)
    except:
        print("Failed to save " + docName)

