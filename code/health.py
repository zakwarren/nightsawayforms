import docx
import header


def health_and_emergency(doc, docName, directory, event):
    """Build the health and emergency contact form for the camp"""
    print("Generating health and emergency contact form...")
    # heading and blurb paragraph
    doc.add_heading("Health and emergency contact form", 4)
    dateBy = input("When does the health form have to be returned by? ")
    blurb = "Please fill in this health and emergency contact information " \
        + "form and return it to a leader by "
    doc.add_paragraph(blurb + str(dateBy) + ".")
    # data protection statement
    p = doc.add_paragraph()
    runner = p.add_run("Data protection")
    runner.bold = True
    dataProtection = "This form is used to collect information about your young person " \
                    + "for the purpose of the event named below, this is to be used by the " \
                    + "Section Leaders only. As part of this form we collect personal data " \
                    + "about your young person, this detail is required so that we can " \
                    + "register them for the event. This form also collects sensitive " \
                    + "(special category) data about your young person, this detail is " \
                    + "required so that we can offer additional support if required and keep " \
                    + "your young person safe whilst in our care. We may share your personal " \
                    + "data in this form with third parties, we do this for event registration. " \
                    + "These third parties are used on the basis that they align with our data " \
                    + "privacy policies. We take your personal data privacy seriously. The data " \
                    + "you provide to us is securely stored (based on local arrangements) and " \
                    + "will be kept for 2 months after the event for any queries that arise " \
                    + "before being securely destroyed. For further detail please visit our " \
                    + "Data Protection Policy "
    r = p.add_run(dataProtection)
    web = "https://scouts.org.uk/media/927472/SCOUTS-data-protection.pdf"
    header.add_hyperlink(p, "here", web)
    p.add_run(". ")

    # personal information section
    p = doc.add_paragraph()
    p.add_run("Name of young person: " + "_" * 40)
    p.add_run(" Date of birth: " + "_" * 20)

    # swimming ability if taking part in water activities
    if event.waterActivities:
        p = doc.add_paragraph()
        p.add_run("Is she/he able to swim 50 metres and stay afloat for " \
            + "five minutes in light clothing? ")
        p.add_run("Yes/No")

    # energency contact details
    p = doc.add_paragraph()
    p.add_run("Emergency contact: " + "_" * 40)
    p.add_run(" Phone: " + "_" * 20)

    # medical details
    med = doc.add_table(rows=4, cols=2)
    # table headings
    cell = med.cell(0, 0)
    cell.width = docx.shared.Cm(10)
    r = cell.paragraphs[0].add_run("Doctor's name and contact details:")
    r.bold = True

    cell = med.cell(0, 1)
    cell.width = docx.shared.Cm(10)
    r = cell.paragraphs[0].add_run("Details of any medications currently being taken:")
    r.bold = True

    cell = med.cell(2, 0)
    r = cell.paragraphs[0].add_run("Details of any disablilities, conditions, allergies, " \
        + "special needs, or cultural needs that might affect this event:")
    r.bold = True

    cell = med.cell(2, 1)
    r = cell.paragraphs[0].add_run("Details of any infectious diseases she/he has been in " \
        + "contact with in the last three weeks:")
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = 2
    r = p.add_run("Please use the back of this form if more space is required")
    r.italic = True

    # consent statement
    p = doc.add_paragraph()
    consent = "If it becomes necessary for the above named young person to receive " \
        + "medical treatment and I cannot be contacted to authorise this, I hereby " \
        + "give my general consent to any necessary medical treatment and authorise " \
        + "the Leader in charge to sign any document required by the hospital authorities."
    r = p.add_run(consent)
    r.italic = True

    p = doc.add_paragraph()
    p.add_run("Signed: " + "_" * 40)
    p.add_run(" Date: " + "_" * 20)
    p = doc.add_paragraph()
    p.add_run("Relationship to young person: " + "_" * 60)

    # disclaimer
    p = doc.add_paragraph()
    disclaim = "Note: The medical profession takes the view that the parent's/carer's " \
        + "consent to medical treatment cannot be delegated. This view is explicit in " \
        + "The Children's Act 1989. Thus, the medical consent forms have no legal status " \
        + "and a doctor or nurse insisting on the consent of a parent/carer to a particular " \
        + "treatment has the right to do so. For this reason we do not recommend that Leaders " \
        + "insist on parents/carers signing the statement above. However, it can be a confort " \
        + "to medical staff to have a general consent in advance from parents/carers or to " \
        + "have a Leader on hand able to sign forms required by medical authorities."
    r = p.add_run(disclaim)
    r.font.size = docx.shared.Pt(6)
    r.font.color.rgb = docx.shared.RGBColor(255, 255, 255)

    # save document
    try:
        docName = 'Health form' + docName
        print("Saving:", docName)
        doc.save(directory + docName)
    except:
        print("Failed to save " + docName)
