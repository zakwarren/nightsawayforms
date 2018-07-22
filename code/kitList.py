import docx, json
import header

def kitList(doc, docName, directory, event):
    """Build the kit list for the camp"""
    print("Generating kit list...")
    # add heading and blurb paragraph
    doc.add_heading('Kit list', 4)
    blurb = 'What to bring: '
    doc.add_paragraph(blurb)

    # pull in kit list from JSON data
    with open('config/kitList.json') as jsonData:
        allKit = json.load(jsonData)
    lists = allKit['kit']

    # set list of categories with calculated responses
    preDefined = [
        "bags",
        "clothing",
        "sleeping",
        "hygiene"
    ]

    for key, values in lists.items():
        # add predefined categories as applicable
        if event.nightsAway > 0 and key in preDefined:
            include = 'y'
        
        # manage special categories
        elif key == "web":
            web = values[0]
        elif key == "exclude":
            exclusions = ', '.join(values)
        elif key == "catering" and event.catering:
            include = 'y'
        elif key == "water activities" and event.waterActivities:
            include = 'y'
        
        # check whether to add each specific section
        else:
            include = input("Include " + key + " section (y/n)? ")
        
        # for each section to include add a bullet point list of kit
        if include == 'y' and key != "exclude" and key != "web":
            for kit in values:
                bullet = doc.add_paragraph(kit)
                bullet.style = 'List Bullet'

    # note any exclusions
    if exclusions:
        excl = doc.add_paragraph()
        excl.add_run("IMPORTANT: Items such as ")
        excl.add_run(exclusions)
        excl.add_run(" are NOT permitted on camp. They’re highly likely to get damaged in the camp environment.")

    # add closing information paragraphs
    close = doc.add_paragraph()
    if web:
        close.add_run("Guidance on kit (particularly if you intend to buy anything) can be found on the ")
        header.add_hyperlink(close, "camp kit page of the website", web)
        close.add_run(". ")
    close.add_run("Feel free to ask the leaders if you have any questions.")

    # save document
    try:
        docName = "Kit list" + docName
        print("Saving:", docName)
        doc.save(directory + docName)
    except:
        print("Failed to save " + docName)
