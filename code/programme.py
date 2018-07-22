import docx

def programme(doc, docName, directory, event):
    """Build the programme for the camp"""
    print("Generating programme...")
    # add heading and blurb paragraph
    doc.add_heading('Programme', 4)
    blurb1 = 'As with all Scouting plans, this programme is designed with a large degree of flexibility. '
    blurb2 = 'Activities and timings are subject to change depending on weather, river conditions, how previous activities go, etc.'
    doc.add_paragraph(blurb1 + blurb2)

    # add main table for the programme
    rowCount = event.nightsAway + 2
    prog = doc.add_table(rows=rowCount, cols=2)
    prog.style = 'Table Grid'
    # add headings
    cell = prog.cell(0, 0)
    cell.width = docx.shared.Cm(3)
    p = cell.add_paragraph()
    r = p.add_run('Day')
    r.bold = True
    cell = prog.cell(0, 1)
    cell.width = docx.shared.Cm(20)
    p = cell.add_paragraph()
    r = p.add_run('Activity')
    r.bold = True

    # populate main table for the programme
    daysAway = event.nightsAway + 1
    i = 1
    while i <= daysAway:
        # add day
        cellD = prog.cell(i, 0)
        d = i - 1
        day = event.daysCamping[d]
        cellD.text = day
        # add programme details
        cellA = prog.cell(i, 1)
        a = 1
        while True:
            act = input("How many activities will there be on " + day + "? ")
            # check is a number
            try:
                act = int(act)
                break
            except:
                print("That's not a number!")
        # get activities and append to programme
        while a <= act:
            activity = input(day + " activity " + str(a) + ": ")
            if a == 1:
                cellA.text = activity
            else:
                pA = cellA.add_paragraph()
                pA.add_run(activity)
            a += 1
        # iterate through loop
        i += 1

    # save document
    try:
        docName = 'Programme' + docName
        print("Saving:", docName)
        doc.save(directory + docName)
    except:
        print("Failed to save " + docName)

