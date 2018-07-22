import docx, json

def riskAssessment(doc, docName, directory, event):
    """Build the risk assessment for the camp"""
    print("Generating risk assessment...")
    # add heading
    doc.add_heading('Risk assessment', 4)
 
    # add main table for the risk assessment
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    # add headings
    cell = table.cell(0, 0)
    cell.width = docx.shared.Cm(5)
    p = cell.add_paragraph()
    r = p.add_run('Risk')
    r.bold = True
    cell = table.cell(0, 1)
    cell.width = docx.shared.Cm(15)
    p = cell.add_paragraph()
    r = p.add_run('Mitigation')
    r.bold = True
    cell = table.cell(0, 2)
    cell.width = docx.shared.Cm(5)
    p = cell.add_paragraph()
    r = p.add_run('Contingency')
    r.bold = True

    # pull in risk data from JSON data
    with open('config/risks.json') as jsonData:
        allRisks = json.load(jsonData)
    risks = allRisks['risks']
    row = 0

    # set list of categories with calculated responses
    preDefined = [
        "health",
        "environmental"
    ]

    for key, values in risks.items():
        # add predefined categories as applicable
        if event.nightsAway > 0 and key in preDefined:
            include = 'y'
        
        # manage special categories
        elif key == "catering" and event.catering:
            include = 'y'
        elif key == "sharps" and event.sharps:
            include = 'y'
        elif key == "water activities" and event.waterActivities:
            include = 'y'
        
        # check whether to add each specific section
        else:
            include = input("Include " + key + " section (y/n)? ")
        
        # for each section to include add a row for assessment
        if include == 'y':
            for assessment in values:
                table.add_row()
                row += 1
                # add risk
                cell = table.cell(row, 0)
                cell.add_paragraph(assessment.capitalize())
                # check mitigation for dynamic components
                mitigation = values[assessment]['mitigation']
                if "[XXX]" in mitigation:
                    print("User input required for risk assessment:\n")
                    print(key.capitalize() + ": " + mitigation, end="\n\n")
                    lines = []
                    i = 0
                    print("Fill in the details to replace [XXX] then press enter. Leave the line blank to finish entering lines.")
                    while True:
                        i += 1
                        line = input("Enter details: ")
                        if line:
                            lines.append(line)
                        else:
                            break
                    newX = '\n'.join(lines)
                    mitigation = mitigation.replace("[XXX]", newX)
                # add mitgation
                cell = table.cell(row, 1)
                cell.add_paragraph(mitigation)
                # check contingency for dynamic components
                contingency = values[assessment]['contingency']
                if "[XXX]" in contingency:
                    print("User input required for risk assessment:\n")
                    print(contingency, end="\n\n")
                    lines = []
                    i = 0
                    print("Fill in the details to replace [XXX] then press enter. Leave the line blank to finish entering lines.")
                    while True:
                        i += 1
                        line = input("Enter details: ")
                        if line:
                            lines.append(line)
                        else:
                            break
                    newX = '\n'.join(lines)
                    contingency = contingency.replace("[XXX]", newX)
                # add contingency
                cell = table.cell(row, 2)
                cell.add_paragraph(contingency)

    # save document
    try:
        docName = 'Risk assessment' + docName
        print("Saving:", docName)
        doc.save(directory + docName)
    except:
        print("Failed to save " + docName)

