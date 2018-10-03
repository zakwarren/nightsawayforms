import docx
import openpyxl
import PIL
from PIL import Image


def resize_logo(logo):
    """resize the logo to the required size while maintaining aspect ratio"""
    img = PIL.Image.open(logo)
    basewidth = 100
    wpercent = (basewidth / float(img.size[0]))
    hsize = int((float(img.size[1]) * float(wpercent)))
    img = img.resize((basewidth, hsize), PIL.Image.ANTIALIAS)
    newImg = 'config/resizedLogo.png'
    img.save(newImg)
    return newImg


def add_hyperlink(paragraph, text, url):
    """Adds a hyperlink to a paragraph"""
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r.append(hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.rgb = docx.shared.RGBColor(0, 149, 200)
    r.font.underline = True

    return hyperlink


def word_heading(myGroup, event):
    """Initialise a document and apply the header"""
    # create document object 
    doc = docx.Document()

    # create header table
    table = doc.add_table(rows=1, cols=2)

    # add title
    cell1 = table.cell(0, 0)
    cell1.width = docx.shared.Cm(25)
    p = cell1.paragraphs[0]
    p.style = 'Title'
    r = p.add_run()
    r.add_text(myGroup.section)
    
    # build sub title
    subTitle = event.camp + ' at ' + event.location + ', from ' + event.dateRange
    # add subtitle 
    p = cell1.add_paragraph()
    r = p.add_run()
    r.add_text(subTitle)
    # add logo
    cell2 = table.cell(0, 1)
    cell2.width = docx.shared.Cm(3)
    p = cell2.paragraphs[0]
    r = p.add_run()
    r.add_picture(myGroup.logo, width=docx.shared.Cm(2.41))

    # generate document name suffix and return all to main programme
    docName = ' - ' + event.camp + ' ' + event.campYear + '.docx'
    return doc, docName

def excel_heading(myGroup, event):
    """Intialise a workbook and apply the header"""
    # create spreadsheet object and capture worksheet
    wb = openpyxl.Workbook()
    ws = wb.worksheets[0]
    ws.title = "equipment"

    # add group
    ws.cell(column=6, row=1).value = myGroup.group
    ws.cell(column=6, row=1).font = openpyxl.styles.Font(size=28)
    ws.cell(column=6, row=1).alignment = openpyxl.styles.Alignment(horizontal="center")

    # resize logo
    smallLogo = resize_logo(myGroup.logo)
    # add logo
    img = openpyxl.drawing.image.Image(smallLogo)
    ws.add_image(img, 'A1')

    # add established year
    ws.cell(column=6, row=4).value = "Established " + str(myGroup.established)
    ws.cell(column=6, row=4).font = openpyxl.styles.Font(size=8)

    # add charity number
    ws.cell(column=1, row=6).value = "Registered charity no. " + str(myGroup.charity)
    ws.cell(column=1, row=6).font = openpyxl.styles.Font(size=8)

    # add RN recognition number if applicable
    if myGroup.rnNumber != 'na':
        ws.cell(column=10, row=6).value = "Admiralty recognised group no. " + str(myGroup.rnNumber)
        ws.cell(column=10, row=6).font = openpyxl.styles.Font(size=8)

    # generate workbook name suffix and return all to main programme
    bookName = ' - ' + event.camp + ' ' + event.campYear + '.xlsx'
    #return workbook, worksheet, bookName
    return wb, ws, bookName

